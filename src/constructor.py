from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import cv2
import os
import uuid
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


import cv2
import numpy as np

def enhance_visual_element(cv_image):

    try:

        lab = cv2.cvtColor(cv_image, cv2.COLOR_BGR2LAB)
        l, a, b = cv2.split(lab)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(4, 4))
        cl = clahe.apply(l)
        limg = cv2.merge((cl, a, b))
        contrast_enhanced = cv2.cvtColor(limg, cv2.COLOR_LAB2BGR)


        denoised = cv2.bilateralFilter(contrast_enhanced, 9, 75, 75)

        kernel = np.array([[-1, -1, -1],
                           [-1,  9, -1],
                           [-1, -1, -1]])
        sharpened = cv2.filter2D(denoised, -1, kernel)
        
        return sharpened

    except Exception as e:
        print(f"Error enhancing visual element: {e}")
        return cv_image # Return the original if enhancement fails


def matrix_word(content_string):

    try : 
        text_matrix = []
        lines = content_string.split('\n')
        for line in lines:
            words = line.split()
            text_matrix.append(words)
        return text_matrix
    except Exception as e :
        print(f"Error in matrix_word: {e}")
        return None


def save_temp_image(cv_image):
    
    temp_folder = "temp_images"
    os.makedirs(temp_folder, exist_ok=True)
    file_path = os.path.join(temp_folder, f"{uuid.uuid4()}.png")
    cv2.imwrite(file_path, cv_image)
    return file_path



def text_constructor(container, element, space_befor_pt = None):
    try :
        content_string = element.get("content", "")
        if not content_string:
            print("Error: No content found")
            return
        text_matrix = matrix_word(content_string)

        props = element.get("font_properties", {})
        text = container.add_paragraph()

        if space_befor_pt and space_befor_pt > 0 :
            text.paragraph_format.space_before = Pt(space_befor_pt)

        alignment_str = element.get("alignment", "left")
        if alignment_str == "center":
            text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment_str == "right":
            text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            text.alignment = WD_ALIGN_PARAGRAPH.LEFT



        for i, line in enumerate(text_matrix):
            for j, word in enumerate(line):
                run = text.add_run(word)

                if props.get("weight") == "bold":
                    run.bold = True

                font_size_str = props.get("size", "medium")
                if font_size_str == "large":
                    run.font.size = Pt(16)

                elif font_size_str == "small":
                    run.font.size = Pt(12)
                else:
                    run.font.size = Pt(14)

                if j < len(line) - 1:
                    text.add_run(" ")


            if i < len(text_matrix) - 1:
                run = text.add_run()
                run.add_break()

    except Exception as e :
        print(f"Error in text_constructor: {e}")


def visual_constructor(doc, element, cv_image, page_width_inches, page_height_inches):
    tmp_path = None
    try :
        img = cv_image
        abs_box = element["absolute_bbox"]
        rel_box = element["relative_bbox"]

        img_h, img_w = cv_image.shape[:2]
        x1 = max(0, abs_box["x_start"])
        y1 = max(0, abs_box["y_start"])
        x2 = min(img_w, abs_box["x_end"])
        y2 = min(img_h, abs_box["y_end"])


        cropped = img[y1:y2, x1:x2]
        enhanced = enhance_visual_element(cropped)

        tmp_path = save_temp_image(enhanced)


        text = doc.add_paragraph()
        text.paragraph_format.space_before = Pt(0)
        text.paragraph_format.space_after = Pt(0)


        alignment_str = element.get("alignment", "left")
        if alignment_str == "center":
            text.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif alignment_str == "right":
            text.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        else:
            text.alignment = WD_ALIGN_PARAGRAPH.LEFT

        relative_width = rel_box["x_end"] - rel_box["x_start"]
        image_width_inches = page_width_inches * relative_width

        relative_height = rel_box["y_end"] - rel_box["y_start"]
        image_height_inches = page_height_inches * relative_height


        run = text.add_run()
        run.add_picture(tmp_path, width=Inches(image_width_inches), height=Inches(image_height_inches))


    except Exception as e :
        print(f"Error in visual_construct : {e}")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)



def table_constructor(container, element, page_width_inches, space_before_pt=None):
    try:
        if space_before_pt and space_before_pt > 0:
            p = container.add_paragraph()
            p.paragraph_format.space_before = Pt(space_before_pt)
            p.paragraph_format.space_after = Pt(0)


        table_data = element.get("table_structure", [])
        if not table_data or not table_data[0]:
            print("table with no structure: Skipping.")
            return
        
        rows = len(table_data)
        cols = len(table_data[0]) if rows > 0 else 0
    
        table = container.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        alignment_str = element.get("alignment", "left")


        if alignment_str == "center":
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

        elif alignment_str == "right":
            table.alignment = WD_TABLE_ALIGNMENT.RIGHT

        else:
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

        
        rel_box = element.get("relative_bbox")
        if rel_box:
            table.autofit = False 
            table.allow_autofit = False            
            relative_width = rel_box["x_end"] - rel_box["x_start"]
            total_table_width = Inches(page_width_inches * relative_width)
            table.width = total_table_width


        for r, row in enumerate(table_data):
            for c, cell in enumerate(row):
                table.cell(r, c).text = cell.get("text", "")

    except Exception as e :
        print(f"failed to make a table construct : {e}")



def docx_constructor(layout_data, image_target, output_path):

    try:
        doc = Document()
        
        section = doc.sections[0]
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        
        usable_page_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
        usable_page_height = section.page_height.inches - section.top_margin.inches - section.bottom_margin.inches


        elements = layout_data.get("elements", [])
        if not elements:
            print("No elements to process.")
            doc.save(output_path)
            return

        elements.sort(key=lambda e: e.get("reading_order_index", 0))
        print("starting document construction")

        for element in elements:
            element_type = element.get("type")

            if element_type in ["header", "paragraph", "footer"]:
                text_constructor(doc, element)
            elif element_type == "table":
                table_constructor(doc, element, usable_page_width)
            elif element_type in ["stamp", "signature", "person_headshot", "signed_headshot", "signed_stamp"]:
                visual_constructor(doc, element, image_target, usable_page_width, usable_page_height)

        doc.save(output_path)
        print(f"Document with precise spacing successfully saved in: {output_path}")

    except Exception as e:
        print(f"failed to make document construct : {e}")
        return None
