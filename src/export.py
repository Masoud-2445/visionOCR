import cv2
from reportlab.pdfgen import canvas
import os
import tempfile
from reportlab.lib.colors import Color
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement
from typing import Dict
from docx.shared import Pt, Inches



def overlay_ocr_to_pdf(image_np, export_data, output_pdf_path):
    """
    Create a PDF with the input image as background and overlay text from OCR data
    with a semi-transparent background, similar to Google Lens Translate.
    
    Args:
        image_np (numpy.ndarray): OpenCV image array
        export_data (dict): Dictionary from OCREngine.prepare_for_export() method
                           Contains 'lines' with text items and 'image_size'
        output_pdf_path (str): Path to save the output PDF
    """
    height, width = image_np.shape[:2]

    width_pt = width * 0.75
    height_pt = height * 0.75

    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_img_file:
        temp_img_path = temp_img_file.name
        cv2.imwrite(temp_img_path, image_np)
    
    try:
        c = canvas.Canvas(output_pdf_path, pagesize=(width_pt, height_pt))
        
        c.drawImage(temp_img_path, 0, 0, width=width_pt, height=height_pt)

        for line in export_data["lines"]:
            line_text = line["text"]

            if not line_text or line_text.isspace():
                continue

            line_items = line["items"]
            if not line_items:
                continue

            min_left = min(item["bbox"][0] for item in line_items)
            min_top = min(item["bbox"][1] for item in line_items)
            max_right = max(item["bbox"][0] + item["bbox"][2] for item in line_items)
            max_bottom = max(item["bbox"][1] + item["bbox"][3] for item in line_items)

            left_px = min_left
            top_px = min_top
            width_px = max_right - min_left
            height_px = max_bottom - min_top

            left = left_px * 0.75
            box_width = width_px * 0.75
            box_height = height_px * 0.75
            top = height_pt - (top_px * 0.75) - box_height
            

            font_size = max(10, min(box_height * 0.7, 72))  # Min 10pt, max 72pt
            c.setFont("Helvetica", font_size)
                
            padding_x = box_width * 0.03
            padding_y = box_height * 0.1

            ext_left = left - padding_x
            ext_top = top - padding_y
            ext_width = box_width + (padding_x * 2)
            ext_height = box_height + (padding_y * 2)
            
            c.saveState()

            c.setFillColor(Color(1, 1, 1, alpha=0.7))
            c.setStrokeColor(Color(0.8, 0.8, 0.8, alpha=0.8))
            c.setLineWidth(0.5)
            
            corner_radius = min(5, ext_height/4)  # Reasonable corner radius
            c.roundRect(ext_left, ext_top, ext_width, ext_height, 
                       corner_radius, stroke=1, fill=1)

            c.setFillColor(Color(0, 0, 0, alpha=1))

            text_x = left + padding_x  # Add some padding from the left edge
            text_y = top + box_height * 0.3  # Adjust baseline position

            c.drawString(text_x, text_y, line_text)

            c.restoreState()

        c.save()   
    finally:
        if os.path.exists(temp_img_path):
            os.remove(temp_img_path)
            
    return output_pdf_path


def export_ocr_to_word(export_data: Dict, output_docx_path: str, 
                                       font_scale_factor: float = 1.3, rtl_support: bool = False) -> None:
    """
    Exports OCR result with PDF-like spacing control.
    
    Args:
        export_data (Dict): A dictionary containing OCR result data.
        output_docx_path (str): Path to save the generated .docx file.
        font_scale_factor (float): Multiplier for base font size (1.3 = 30% larger).
        rtl_support (bool): Whether to apply right-to-left formatting.
    """
    try:
        doc = Document()

        # Set narrow margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Scale the base font size
        scaled_font_size = int(11 * font_scale_factor)
        print(f"Using scaled font size: {scaled_font_size}pt")

        lines = export_data.get("lines", [])

        # Sort lines by top bounding box coordinate
        if lines:
            lines_with_top = []
            for line in lines:
                if line.get("items") and line["items"][0].get("bbox"):
                    top_pos = line["items"][0]["bbox"][1]
                    lines_with_top.append((top_pos, line))
                else:
                    lines_with_top.append((0, line))
            lines_with_top.sort(key=lambda x: x[0])
            sorted_lines = [line for _, line in lines_with_top]
        else:
            sorted_lines = lines

        for i, line in enumerate(sorted_lines):
            line_text = line.get("text", "")
            items = line.get("items", [])

            if not line_text.strip():
                continue

            paragraph = doc.add_paragraph()

            if rtl_support:
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                pPr = paragraph._element.get_or_add_pPr()
                bidi = OxmlElement("w:bidi")
                pPr.append(bidi)

            # Calculate line dimensions (similar to PDF box calculations)
            if items:
                line_heights = []
                for item in items:
                    if item.get("bbox"):
                        height = item["bbox"][3]
                        line_heights.append(height)
                
                if line_heights:
                    avg_line_height = sum(line_heights) / len(line_heights)
                    # Apply similar padding concept as PDF (padding_y = box_height * 0.1)
                    vertical_padding = avg_line_height * 0.1 * 0.75  # Convert to points
                    
                    # Set spacing based on line height (similar to PDF approach)
                    space_before = max(0, vertical_padding * 0.5)
                    space_after = max(2, vertical_padding * 0.8)
                    
                    # Line spacing based on text height
                    line_spacing = 1.0 + (avg_line_height * 0.001)  # Subtle adjustment
                else:
                    space_before = 0
                    space_after = 2
                    line_spacing = 1.1
            else:
                space_before = 0
                space_after = 2
                line_spacing = 1.1

            paragraph.paragraph_format.space_before = Pt(space_before)
            paragraph.paragraph_format.space_after = Pt(space_after)
            paragraph.paragraph_format.line_spacing = line_spacing

            if items:
                for idx, item in enumerate(items):
                    item_text = item.get("text", "")
                    if not item_text.strip():
                        continue

                    run = paragraph.add_run(item_text)
                    run.font.size = Pt(scaled_font_size)
                    run.font.bold = False
                    run.font.italic = False
                    run.font.underline = False

                    if idx < len(items) - 1:
                        space = paragraph.add_run(" ")
                        space.font.size = Pt(scaled_font_size)
                        space.font.bold = False
            else:
                run = paragraph.add_run(line_text)
                run.font.size = Pt(scaled_font_size)
                run.font.bold = False
                run.font.italic = False
                run.font.underline = False

        doc.save(output_docx_path)
        print(f"Word document saved to: {output_docx_path}")

    except Exception as e:
        print("Error exporting OCR result to Word:")
        print(f"Exception: {e}")







        # development stage ---------------------------------------------